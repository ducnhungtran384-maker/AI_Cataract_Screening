import re
import os
import keyword
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name_en='Times New Roman', font_name_cn='宋体', size=None, bold=False, italic=False):
    if not font_name_cn:
        # Default fallback if not specified
        font_name_cn = '宋体'

    run.font.name = font_name_en

    r = run._element
    rPr = r.get_or_add_rPr()
    
    # 强制设置所有语种的字体 (Use Localized Names: 宋体 instead of SimSun)
    fonts = rPr.get_or_add_rFonts()
    fonts.set(qn('w:eastAsia'), font_name_cn)
    fonts.set(qn('w:ascii'), font_name_en)
    fonts.set(qn('w:hAnsi'), font_name_en)
    fonts.set(qn('w:cs'), font_name_en)
    
    # 重点：告诉 Word 这是一个东亚语言的 Run，优先使用 EastAsia 字体
    fonts.set(qn('w:hint'), 'eastAsia')

    if size:
        run.font.size = size
    if bold:
        run.bold = True
    if italic:
        run.italic = True

def parse_markdown_to_docx(md_file_path, docx_file_path):
    doc = Document()
    
    # 1. Setup Styles (Default Normal)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12) 
    
    # Deep XML Hack for Style Font
    rPr = style._element.get_or_add_rPr()
    fonts = rPr.get_or_add_rFonts()
    fonts.set(qn('w:eastAsia'), '宋体')
    fonts.set(qn('w:ascii'), 'Times New Roman')
    
    style.paragraph_format.space_after = Pt(0) 
    style.paragraph_format.line_spacing = 1.25

    # 2. Heading Styles
    for i in range(1, 4):
        try:
            h_style = doc.styles[f'Heading {i}']
            h_style.font.name = 'Arial'
            
            # Deep XML Hack for Headings
            rPr = h_style._element.get_or_add_rPr()
            fonts = rPr.get_or_add_rFonts()
            fonts.set(qn('w:eastAsia'), '黑体') # SimHei
            fonts.set(qn('w:ascii'), 'Arial')
            
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            if i == 1:
                h_style.font.size = Pt(16)
                h_style.font.bold = True
                h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 2:
                h_style.font.size = Pt(14)
                h_style.font.bold = True
            elif i == 3:
                h_style.font.size = Pt(13)
                h_style.font.bold = True
        except:
            pass

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []
    in_code_block = False
    code_block_content = []

    # Track state
    skip_toc = False
    
    for line in lines:
        stripped_line = line.strip()
        
        # --- Code Blocks ---
        # --- Code Blocks with Highlighting ---
        if stripped_line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_content = []
            else:
                in_code_block = False
                if code_block_content:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing' 
                    p.paragraph_format.left_indent = Pt(24) # Indent code block
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.space_before = Pt(6)
                    
                    # 1. Background Shading (Black)
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '000000') # Black background
                    p._p.get_or_add_pPr().append(shd)

                    for code_line in code_block_content:
                         # 2. Simple Tokenizer for Coloring
                        tokens = re.split(r'(\s+|[().,\[\]:;\'"=\+\-\*/%{}<>])', code_line) 
                        
                        # Check for full line comment first (simpler)
                        if code_line.strip().startswith('#'):
                             run = p.add_run(code_line + '\n')
                             run.font.name = 'Consolas'
                             run.font.size = Pt(9.5)
                             run.font.color.rgb = RGBColor(87, 166, 74) # VS Code Green
                             continue

                        for token in tokens:
                            if not token: continue
                            
                            run = p.add_run(token)
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9.5)
                            
                            # Default Color: White
                            run.font.color.rgb = RGBColor(220, 220, 220) # Light Grey/White
                            
                            clean_token = token.strip()
                            # Keyword Coloring 
                            if clean_token in keyword.kwlist:
                                 run.font.color.rgb = RGBColor(86, 156, 214) # VS Code Blue
                            elif clean_token in ['print', 'range', 'len', 'open', 'str', 'int', 'list', 'dict']: 
                                 run.font.color.rgb = RGBColor(220, 220, 170) # Light Yellow
                            elif clean_token in ['self', 'True', 'False', 'None']:
                                 run.font.color.rgb = RGBColor(86, 156, 214) # Blue
                            elif token.startswith("'") or token.startswith('"'):
                                 run.font.color.rgb = RGBColor(214, 157, 133) # VS Code string orange
                            elif token.startswith('#'):
                                 run.font.color.rgb = RGBColor(87, 166, 74) # Green
                            
                            # Enforce EastAsian Font for Code too (comments!)
                            set_font(run, font_name_en='Consolas', font_name_cn='宋体', size=Pt(9.5))
                        
                        # Add newline
                        p.add_run('\n')

                code_block_content = []
            continue
        
        if in_code_block:
            # Preserve empty lines and indentation
            code_block_content.append(line.rstrip()) 
            continue
        
        if not stripped_line:
            continue

        # --- TOC Skipping Logic ---
        # If we see the TOC header or strict markdown links that look like TOC items
        if stripped_line == "# 目  录" or stripped_line == "# 目录":
            # REVERT TO AUTO TOC (User Request: "Add Page Numbers")
            # Static TOC cannot have page numbers. We must use the Field Code.
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            
            # <w:fldChar w:fldCharType="begin"/>
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            # <w:instrText xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
            
            # <w:fldChar w:fldCharType="separate"/>
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            
            # Instruction Text Run (Visible if not updated)
            r_instr = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            color = OxmlElement('w:color')
            color.set(qn('w:val'), 'DC143C') 
            rPr.append(color)
            
            bold = OxmlElement('w:b')
            rPr.append(bold)
            
            r_instr.append(rPr)
            
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = "【目录生成区：请在此处 点击右键 -> 更新域 -> 更新整个目录 (以显示页码)】"
            r_instr.append(t)
            
            # <w:fldChar w:fldCharType="end"/>
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            
            r_element = run._r
            r_element.append(fldChar1)
            r_element.append(instrText)
            r_element.append(fldChar2)
            r_element.append(r_instr) 
            r_element.append(fldChar3)
            
            doc.add_page_break()
            skip_toc = True
            continue

        if skip_toc:
            # If line is a link like [一、...](#...), skip it
            if re.match(r'\[.*\]\(#.*\)', stripped_line):
                continue
            # If we hit a real header, stop skipping
            if stripped_line.startswith('# '):
                skip_toc = False
                # If it's the Main Title
                if "实践课题报告" in stripped_line:
                     doc.add_page_break() # Start fresh after TOC assumption

        # --- Main Title Handling ---
        # Detect the specific main title
        if "《人工智能导论" in stripped_line and stripped_line.startswith('# '):
             text = stripped_line.replace('# ', '').strip()
             p = doc.add_heading(text, level=1)
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             for run in p.runs:
                 run.font.size = Pt(22) # 二号
                 run.font.name = 'SimHei'
                 run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                 run.font.color.rgb = RGBColor(0, 0, 0)
             continue

        # --- Blockquotes (Images inside quotes or Text quotes) ---
        is_blockquote = False
        if stripped_line.startswith('>'):
            is_blockquote = True
            stripped_line = stripped_line.lstrip('>').strip()
            if not stripped_line:
                continue

        # --- Images (Now handles those inside qutoes too) ---
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped_line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            img_path = img_match.group(2)
            
            # Handle file:// URIs
            if img_path.startswith('file:///'):
                img_path = img_path[8:]
            elif img_path.startswith('file://'):
                img_path = img_path[7:]
            
            # Clean up path (remove title if present)
            img_path = img_path.split(' ')[0]
            
            if os.path.isabs(img_path):
                abs_img_path = img_path
            else:
                base_dir = os.path.dirname(md_file_path)
                if img_path.startswith('./'): img_path = img_path[2:]
                abs_img_path = os.path.join(base_dir, img_path)
            
            if os.path.exists(abs_img_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(abs_img_path, width=Inches(5.0)) # A bit smaller to fit
                    
                    # FILTER OUT "1.00" CAPTIONS
                    if alt_text and alt_text != "1.00":
                        # Caption
                        caption_p = doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.paragraph_format.space_after = Pt(12)
                        caption_run = caption_p.add_run(alt_text)
                        set_font(caption_run, size=Pt(10.5), font_name_cn='楷体', italic=True)
                except Exception as e:
                    print(f"Image insert error: {e}")
                    doc.add_paragraph(f"[图片加载失败: {img_path}]")
            else:
                 doc.add_paragraph(f"[图片文件未找到: {img_path}]")
            continue
            
        # If it was a blockquote but NOT an image, it's a text quote/caption
        if is_blockquote:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Captions usually centered or left with indent? 
            # User screenshot shows left aligned block text. Let's make it look like a caption or note.
            # Actually, standard blockquotes should differ.
            # But in this report, > seems used for Figure Captions mostly.
            if "图" in stripped_line or "图注" in stripped_line:
                 p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                 p.paragraph_format.left_indent = Pt(24) # Indent quote
            
            add_run_with_formatting(p, stripped_line)
            # Set style for the whole paragraph run
            for run in p.runs:
                 set_font(run, font_name_cn='楷体', size=Pt(10.5))
            continue

        # --- Headers ---
        header_match = re.match(r'^(#+)\s+(.*)', stripped_line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            text = text.replace('**', '').replace('*', '') 
            
            # Special handling for "Cover Page" header -> hide it or make generic
            if "封面页" in text:
                 # Ensure page break before new major sections if not detected earlier
                 if "一、" in text and level == 1:
                     doc.add_page_break()
            
            # Special handling for "Cover Page" header -> hide it or make generic
            if "封面页" in text:
                 # Ensure page break before new major sections if not detected earlier
                 if "一、" in text and level == 1:
                     doc.add_page_break()
            
            p = doc.add_heading(smart_replace(text), level=min(level, 9))
            # Center alignment for H1
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # --- Lists ---
        if stripped_line.startswith('* ') or stripped_line.startswith('- '):
            text = stripped_line[2:]
            p = doc.add_paragraph()
            p.style = 'List Paragraph' # Use native list style
            add_run_with_formatting(p, text)
            continue
            
        if re.match(r'^\d+\.\s+', stripped_line):
            text = re.sub(r'^\d+\.\s+', '', stripped_line)
            p = doc.add_paragraph()
            p.style = 'List Paragraph'
            # Manual Numbering prefix since robust auto-numbering is hard in python-docx simple pass
            # But 'List Paragraph' adds indent. Let's just prepend the number.
            num_prefix = re.match(r'^\d+\.\s+', stripped_line).group(0)
            add_run_with_formatting(p, num_prefix + text)
            continue
        

        # --- Tables ---
        if stripped_line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            if re.match(r'\|[\s:\-]+\|', stripped_line):
                continue   
            cells = [cell.strip() for cell in stripped_line.strip('|').split('|')]
            table_data.append(cells)
            continue
        else:
            if in_table:
                if table_data:
                    num_cols = len(table_data[0])
                    rows = len(table_data)
                    table = doc.add_table(rows=rows, cols=num_cols)
                    table.style = 'Table Grid'
                    table.autofit = False 
                    for i, row_data in enumerate(table_data):
                        for j, val in enumerate(row_data):
                            if j < num_cols:
                                cell = table.cell(i, j)
                                cell.paragraphs[0].clear() # clear default
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                add_run_with_formatting(p, val)
                in_table = False
                table_data = []
        
        # --- Normal Paragraph ---
        p = doc.add_paragraph()
        
        # Cover Page Field Formatting
        if "课题名称：" in stripped_line or "学期：" in stripped_line or "小组成员" in stripped_line or "临床医学" in stripped_line:
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             p.paragraph_format.line_spacing = 1.5
             add_run_with_formatting(p, stripped_line)
        else:
            p.paragraph_format.first_line_indent = Pt(24) # Indent standard paragraphs (2 chars)
            add_run_with_formatting(p, stripped_line)

    # 3. Add Page Numbers
    add_page_numbers(doc)
    
    # 4. Global Font Force (FINAL HAMMER)
    force_cjk_font_global(doc)
    
    # 5. Set Update Fields (Auto TOC Prompt)
    set_updatefields_true(doc)

    doc.save(docx_file_path)
    print(f"Successfully created {docx_file_path}")

def add_page_numbers(doc):
    """
    Adds page numbers to the footer of every section.
    """
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] # Use default empty paragraph
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        
        # XML Magic for "PAGE" field
        # <w:fldChar w:fldCharType="begin"/>
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        # <w:instrText xml:space="preserve"> PAGE </w:instrText>
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        # <w:fldChar w:fldCharType="separate"/>
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # <w:fldChar w:fldCharType="end"/>
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        r_element = run._r
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)

def add_hyperlink(paragraph, url, text):
    """
    Adds a hyperlink to a paragraph using OXML.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Join the new_run and rPr
    new_run.append(rPr)
    new_run.text = smart_replace(text)
    
    # Style the run
    # Format color blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    # Format underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    # Apply Font
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), "Times New Roman")
    fonts.set(qn('w:eastAsia'), "宋体")
    fonts.set(qn('w:hint'), "eastAsia")
    rPr.append(fonts)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def add_run_with_formatting(paragraph, text):
    """
    Parses **bold**, *italic*, `code`, and [link](url).
    Strips the markdown characters and applies styles.
    Aggressively removes ` characters.
    """
    pattern = r'(\*\*.*?\*\*)|(\*.*?\*)|(`.*?`)|(\[.*?\]\(.*?\))'
    
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
            
        # Bold
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = paragraph.add_run(smart_replace(content.replace('`', '')))
            set_font(run, bold=True)
            
        # Italic
        elif part.startswith('*') and part.endswith('*'):
            content = part[1:-1]
            run = paragraph.add_run(smart_replace(content.replace('`', '')))
            set_font(run, italic=True)
            
        # Code
        elif part.startswith('`') and part.endswith('`'):
            content = part[1:-1]
            run = paragraph.add_run(content) # already stripped wrapper
            run.font.name = 'Consolas'
            
        # Link
        elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            m = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if m:
                link_text = m.group(1)
                link_url = m.group(2)
                # Apply REAL hyperlink
                add_hyperlink(paragraph, link_url, link_text)
            else:
                paragraph.add_run(smart_replace(part.replace('`', '')))

        # Auto-link <url>
        elif part.startswith('<') and part.endswith('>') and 'http' in part:
             url = part[1:-1]
             add_hyperlink(paragraph, url, url)
                
        else:
            # Plain text
            # Fix escapes AND remove <br /> tags
            clean_part = part.replace('`', '').replace('\\_', '_').replace('\\', '') 
            clean_part = re.sub(r'<br\s*/?>', '', clean_part, flags=re.IGNORECASE)
            
            run = paragraph.add_run(smart_replace(clean_part))
            set_font(run)

def smart_replace(text):
    """
    Converts half-width parens to full-width for Chinese context.
    Also removes markdown escape backslashes '\\_' -> '_'.
    """
    return text.replace('(', '（').replace(')', '）').replace('\\_', '_').replace('\\', '')

def force_cjk_font_global(doc):
    """
    Final pass to ensure NO MS Gothic remains. 
    Iterates all paragraphs and runs to enforce SimSun/SimHei.
    User Request: Change Heiti to Songti (SimSun) for better aesthetics.
    """
    for p in doc.paragraphs:
        # Determine if it is a heading
        # User requested Songti even for headers usually, just Bold Songti
        target_font = '宋体' 
        
        for run in p.runs:
            r = run._element
            rPr = r.get_or_add_rPr()
            fonts = rPr.get_or_add_rFonts()
            fonts.set(qn('w:eastAsia'), target_font)
            
            # Keep ascii font as is (likely Times New Roman or Arial set previously)
            if not fonts.get(qn('w:ascii')):
                 fonts.set(qn('w:ascii'), 'Times New Roman')

def set_updatefields_true(doc):
    """
    Forces Word to prompt "Update Fields?" on open, or auto-update TOC.
    This is the best we can do since Python doesn't know page numbers.
    """
    element = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    element.append(update_fields)

if __name__ == "__main__":
    md_path = r"C:\Users\weirui\Desktop\AI_Test\AI辅助白内障筛查实践报告_converted.md"
    docx_path = r"C:\Users\weirui\Desktop\AI_Test\AI辅助白内障筛查实践报告_完美终稿.docx"
    
    try:
        parse_markdown_to_docx(md_path, docx_path)
    except Exception as e:
        print(f"Error: {e}")
